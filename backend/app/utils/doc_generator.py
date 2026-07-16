import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_from_json(structured_contract: dict, output_path: str) -> None:
    """
    Generates a professionally formatted Microsoft Word (.docx) document
    from structured contract JSON data. Contains no AI logic.
    """
    doc = Document()

    # Configure page margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set document default style font to Times New Roman, size 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 1. Document Title
    title_text = structured_contract.get("title", "Freelance Services Agreement").strip()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(24)
    
    title_run = title_p.add_run(title_text)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(18)
    title_run.bold = True

    # 2. Parties Section
    parties = structured_contract.get("parties", {})
    if parties:
        h_parties = doc.add_heading("PARTIES", level=1)
        h_parties.paragraph_format.space_before = Pt(18)
        h_parties.paragraph_format.space_after = Pt(6)
        h_parties.keep_with_next = True
        for run in h_parties.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = None  # Use default black font color

        p_parties = doc.add_paragraph()
        p_parties.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_parties.paragraph_format.line_spacing = 1.15
        p_parties.paragraph_format.space_after = Pt(12)
        
        client_name = parties.get("client", "").strip() or "Client"
        freelancer_name = parties.get("freelancer", "").strip() or "Freelancer"
        
        text_parties = (
            f"This Agreement is entered into by and between the following parties:\n\n"
            f"Client:\n{client_name}\n\n"
            f"Freelancer:\n{freelancer_name}"
        )
        run_parties = p_parties.add_run(text_parties)
        run_parties.font.name = 'Times New Roman'
        run_parties.font.size = Pt(12)

    # 3. Main Sections
    sections = structured_contract.get("sections", [])
    for idx, sec in enumerate(sections):
        heading_text = sec.get("heading", f"Section {idx + 1}").strip().upper()
        content_text = sec.get("content", "").strip()

        # Add section heading
        h_sec = doc.add_heading(heading_text, level=1)
        h_sec.paragraph_format.space_before = Pt(18)
        h_sec.paragraph_format.space_after = Pt(6)
        h_sec.keep_with_next = True
        for run in h_sec.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = None

        # Add content paragraphs (retaining spacing/newlines)
        paragraphs_text = content_text.split("\n\n")
        for para_text in paragraphs_text:
            cleaned_text = para_text.strip()
            if not cleaned_text:
                continue
                
            p_content = doc.add_paragraph()
            p_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_content.paragraph_format.line_spacing = 1.15
            p_content.paragraph_format.space_after = Pt(12)
            
            run_content = p_content.add_run(cleaned_text)
            run_content.font.name = 'Times New Roman'
            run_content.font.size = Pt(12)

    # 4. Signatures Section
    sig_heading = doc.add_heading("SIGNATURES", level=1)
    sig_heading.paragraph_format.space_before = Pt(36)
    sig_heading.paragraph_format.space_after = Pt(18)
    sig_heading.keep_with_next = True
    for run in sig_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = None

    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    
    # Format signature placeholders
    client_name = parties.get("client", "Client").strip()
    freelancer_name = parties.get("freelancer", "Freelancer").strip()

    # Column 0: Client signature block
    cell_client_sig = table.cell(0, 0)
    p_client_sig = cell_client_sig.paragraphs[0]
    p_client_sig.paragraph_format.space_after = Pt(36)  # space to sign
    run_cs1 = p_client_sig.add_run("_________________________________________\n")
    run_cs1.font.name = 'Times New Roman'
    run_cs1.font.size = Pt(12)
    
    cell_client_title = table.cell(1, 0)
    p_client_title = cell_client_title.paragraphs[0]
    run_cs2 = p_client_title.add_run(f"For Client: {client_name}\nDate: _________________")
    run_cs2.font.name = 'Times New Roman'
    run_cs2.font.size = Pt(12)

    # Column 1: Freelancer signature block
    cell_free_sig = table.cell(0, 1)
    p_free_sig = cell_free_sig.paragraphs[0]
    p_free_sig.paragraph_format.space_after = Pt(36)  # space to sign
    run_fs1 = p_free_sig.add_run("_________________________________________\n")
    run_fs1.font.name = 'Times New Roman'
    run_fs1.font.size = Pt(12)
    
    cell_free_title = table.cell(1, 1)
    p_free_title = cell_free_title.paragraphs[0]
    run_fs2 = p_free_title.add_run(f"For Freelancer: {freelancer_name}\nDate: _________________")
    run_fs2.font.name = 'Times New Roman'
    run_fs2.font.size = Pt(12)

    # Ensure widths are consistent (around 3.25 inches per signature block)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.25)

    # Save to the desired output path
    # Ensure parent directory exists
    parent_dir = os.path.dirname(output_path)
    if parent_dir:
        os.makedirs(parent_dir, exist_ok=True)
        
    doc.save(output_path)
